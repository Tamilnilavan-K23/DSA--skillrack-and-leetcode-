from docx import Document

# Initialize document
doc = Document()
doc.add_heading('LeetCode Array Problem Set (Easy to Medium)', 0)

# Section definitions
sections = {
    "Easy Level Array Problems": [
        ("Two Sum", "https://leetcode.com/problems/two-sum/"),
        ("Best Time to Buy and Sell Stock", "https://leetcode.com/problems/best-time-to-buy-and-sell-stock/"),
        ("Remove Duplicates from Sorted Array", "https://leetcode.com/problems/remove-duplicates-from-sorted-array/"),
        ("Contains Duplicate", "https://leetcode.com/problems/contains-duplicate/"),
        ("Merge Sorted Array", "https://leetcode.com/problems/merge-sorted-array/"),
        ("Intersection of Two Arrays II", "https://leetcode.com/problems/intersection-of-two-arrays-ii/"),
        ("Plus One", "https://leetcode.com/problems/plus-one/"),
        ("Move Zeroes", "https://leetcode.com/problems/move-zeroes/"),
        ("Valid Mountain Array", "https://leetcode.com/problems/valid-mountain-array/"),
        ("Find Numbers with Even Number of Digits", "https://leetcode.com/problems/find-numbers-with-even-number-of-digits/")
    ],
    "Medium Level - Two Pointer / Sliding Window": [
        ("3Sum", "https://leetcode.com/problems/3sum/"),
        ("Container With Most Water", "https://leetcode.com/problems/container-with-most-water/"),
        ("Longest Substring Without Repeating Characters", "https://leetcode.com/problems/longest-substring-without-repeating-characters/"),
        ("Minimum Size Subarray Sum", "https://leetcode.com/problems/minimum-size-subarray-sum/"),
        ("Subarray Product Less Than K", "https://leetcode.com/problems/subarray-product-less-than-k/"),
        ("Longest Repeating Character Replacement", "https://leetcode.com/problems/longest-repeating-character-replacement/"),
        ("Fruits into Baskets", "https://leetcode.com/problems/fruit-into-baskets/"),
        ("Permutation in String", "https://leetcode.com/problems/permutation-in-string/"),
        ("Find All Anagrams in a String", "https://leetcode.com/problems/find-all-anagrams-in-a-string/"),
        ("Maximum Points You Can Obtain from Cards", "https://leetcode.com/problems/maximum-points-you-can-obtain-from-cards/")
    ],
    "Medium Level - Prefix Sum / Hashing / Binary Search": [
        ("Subarray Sum Equals K", "https://leetcode.com/problems/subarray-sum-equals-k/"),
        ("Maximum Size Subarray Sum Equals k", "https://leetcode.com/problems/maximum-size-subarray-sum-equals-k/"),
        ("Range Sum Query - Immutable", "https://leetcode.com/problems/range-sum-query-immutable/"),
        ("Find First and Last Position of Element in Sorted Array", "https://leetcode.com/problems/find-first-and-last-position-of-element-in-sorted-array/"),
        ("Search in Rotated Sorted Array", "https://leetcode.com/problems/search-in-rotated-sorted-array/"),
        ("Peak Index in a Mountain Array", "https://leetcode.com/problems/peak-index-in-a-mountain-array/"),
        ("Find Pivot Index", "https://leetcode.com/problems/find-pivot-index/"),
        ("Matrix Cells in Distance Order", "https://leetcode.com/problems/matrix-cells-in-distance-order/"),
        ("K-diff Pairs in an Array", "https://leetcode.com/problems/k-diff-pairs-in-an-array/"),
        ("Sum of Subarray Minimums", "https://leetcode.com/problems/sum-of-subarray-minimums/")
    ]
}

# Add sections to document
for title, problems in sections.items():
    doc.add_heading(title, level=1)
    for i, (name, url) in enumerate(problems, 1):
        doc.add_paragraph(f"{i}. {name}\n   {url}", style='List Number')

# Save document
doc.save("LeetCode_Array_Problems.docx")
